"""Professional DOCX exporter using python-docx for Google Docs and Word compatibility."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class DocxExporter:
    def __init__(self, project_root: Path):
        self.project_root = project_root
        self.chapters_dir = project_root / "chapters"
        self.exports_dir = project_root / "exports" / "gdoc"

    def export(self) -> Path:
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()
        self._setup_styles(doc)
        self._setup_page(doc)
        self._add_toc(doc)

        files = self._get_final_files()
        if not files:
            raise RuntimeError("No final chapter files found.")

        first_heading = True
        for filepath in files:
            content = Path(filepath).read_text(encoding="utf-8")
            content = self._strip_front_matter(content)
            self._add_markdown_content(doc, content, first_heading)
            first_heading = False

        output = self.exports_dir / "mastering-legal-transcription.docx"
        doc.save(str(output))
        return output

    def _setup_styles(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Georgia"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        for level, size in [(1, 18), (2, 14), (3, 12)]:
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Georgia"
            heading_style.font.size = Pt(size)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

        if "Callout" not in doc.styles:
            callout_style = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
            callout_style.font.name = "Georgia"
            callout_style.font.size = Pt(10.5)
            callout_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _setup_page(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _get_final_files(self) -> list[str]:
        files = []
        for chapter_dir in sorted(self.chapters_dir.glob("[0-9][0-9]-*")):
            final_dir = chapter_dir / "final"
            if final_dir.exists():
                for md in final_dir.glob("*.md"):
                    files.append(str(md))
        return files

    def _strip_front_matter(self, text: str) -> str:
        if text.startswith("---"):
            parts = text.split("---", 2)
            if len(parts) == 3:
                return parts[2].lstrip()
        return text

    def _add_toc(self, doc: Document) -> None:
        p = doc.add_paragraph()
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        run._r.append(fld)
        doc.add_page_break()

    def _add_markdown_content(self, doc: Document, content: str, first_heading: bool) -> None:
        lines = content.split("\n")
        in_code_block = False
        code_lines: list[str] = []
        table_buffer: list[str] = []

        for line in lines:
            if line.strip().startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9.5)
                    p.paragraph_format.left_indent = Inches(0.5)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if line.strip().startswith("|") and line.strip().endswith("|"):
                table_buffer.append(line)
                continue
            if table_buffer:
                self._add_table(doc, table_buffer)
                table_buffer = []

            if re.match(r"^\\[CALLOUT:", line.strip(), re.IGNORECASE):
                p = doc.add_paragraph(style="Callout")
                self._add_formatted_text(p, line.strip())
                continue

            if re.match(r"^\\s*[-*]\\s+", line):
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(p, re.sub(r"^\\s*[-*]\\s+", "", line))
                continue
            if re.match(r"^\\s*\\d+\\.\\s+", line):
                p = doc.add_paragraph(style="List Number")
                self._add_formatted_text(p, re.sub(r"^\\s*\\d+\\.\\s+", "", line))
                continue

            if line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=3)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                if not first_heading:
                    doc.add_page_break()
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("# "):
                if not first_heading:
                    doc.add_page_break()
                doc.add_heading(line[2:].strip(), level=1)
            elif line.strip() == "" or line.strip() == "---":
                continue
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

        if table_buffer:
            self._add_table(doc, table_buffer)

    def _add_table(self, doc: Document, lines: list[str]) -> None:
        rows = []
        for line in lines:
            if re.match(r"^\\|\\s*-", line):
                continue
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(cells)
        if not rows:
            return
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        for idx, cell in enumerate(rows[0]):
            table.rows[0].cells[idx].text = cell
        for row in rows[1:]:
            cells = table.add_row().cells
            for idx, cell in enumerate(row):
                cells[idx].text = cell

    def _add_formatted_text(self, paragraph, text: str) -> None:
        parts = re.split(r"(\\*\\*.*?\\*\\*|\\*.*?\\*|`.*?`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
            else:
                paragraph.add_run(part)
